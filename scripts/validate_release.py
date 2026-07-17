from __future__ import annotations

import hashlib
import json
import py_compile
import sys
from pathlib import Path

import joblib
import fitz
import nbformat
import numpy as np
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from stackoverflow_clustering.reporting import REPORT_DETAILS


def _json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while block := handle.read(1024 * 1024):
            digest.update(block)
    return digest.hexdigest()


def _run_complex_script_font(run) -> str | None:
    """Return the explicit complex-script font stored on a run, if any."""
    properties = run._element.rPr
    if properties is None or properties.rFonts is None:
        return None
    return properties.rFonts.get(qn("w:cs"))


def _assert_heading_fonts(document: Document, expected: str) -> None:
    """Validate rendered heading runs, even after Word normalizes styles to themes."""
    headings = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style is not None and paragraph.style.name.startswith("Heading")
    ]
    assert headings, "Report contains no heading paragraphs"
    for paragraph in headings:
        visible_runs = [run for run in paragraph.runs if run.text.strip()]
        assert visible_runs, f"Heading has no visible run: {paragraph.text!r}"
        assert all(
            _run_complex_script_font(run) == expected for run in visible_runs
        ), f"Unexpected heading font: {paragraph.text!r}"


def validate() -> dict:
    artifacts = ROOT / "artifacts"
    processed = ROOT / "data" / "processed"
    reports = ROOT / "reports"
    tables = reports / "tables"

    phase1 = _json(artifacts / "phase1_run_summary.json")
    phase2 = _json(artifacts / "phase2_run_summary.json")
    phase3 = _json(artifacts / "phase3_run_summary.json")
    bonus = _json(artifacts / "bonus_summary.json")
    registry = _json(artifacts / "model_registry.json")
    assert phase1["status"] == phase2["status"] == phase3["status"] == "completed"
    assert phase1["ingestion"]["rows"] == 65_437
    assert phase1["cleaning"]["rows_cleaned"] == 60_023
    assert phase1["features"]["representations"]["X_full"] == [60_023, 328]

    phase2_scores = pd.read_csv(tables / "phase2_model_family_scores.csv")
    assert len(phase2_scores) == phase2["model_families"]["configurations_evaluated"] == 78
    assert phase2_scores["family"].nunique() == 5
    bic_row = phase2_scores.loc[phase2_scores["bic"].idxmin()]
    expected_gmm = f"{bic_row['algorithm']}_k{int(bic_row['k_requested'])}"
    assert expected_gmm == phase2["model_families"]["gmm_bic_winner"]

    labels = np.load(processed / "phase3_consensus_labels.npy")
    centroids = np.load(processed / "phase3_consensus_centroids.npy")
    profiles = pd.read_csv(tables / "phase3_cluster_profiles.csv")
    assert len(labels) == 60_023
    assert centroids.shape == (phase3["consensus_k"], 50)
    assert profiles["respondents"].sum() == len(labels)
    assert phase3["cluster_sizes"] == {
        str(int(cluster)): int(count) for cluster, count in zip(*np.unique(labels, return_counts=True))
    }
    assert 0 <= phase3["validation_accuracy"] <= 1
    assert 0 <= phase3["explanation_tree_validation_accuracy"] <= 1

    assert phase3["bonus_status"] == bonus["status"] == "completed"
    assert phase3["bonus_points_targeted"] == bonus["points_targeted"] == 15
    nmf_labels = np.load(processed / "bonus_nmf_labels.npy")
    assert len(nmf_labels) == len(labels)
    assert len(np.unique(nmf_labels)) == bonus["nmf"]["selected_components"]
    nmf_model = artifacts / "models" / "bonus_nmf_technology.joblib"
    joblib.load(nmf_model)
    assert _sha256(nmf_model) == bonus["nmf"]["model_sha256"]
    confidence = pd.read_csv(tables / "bonus_metric_confidence_intervals.csv")
    assert set(confidence["metric"]) == {
        "silhouette",
        "calinski_harabasz_per_row",
        "davies_bouldin",
        "ari_vs_consensus",
        "nmi_vs_consensus",
    }
    assert {"consensus", "nmf_high_dimensional"}.issubset(set(confidence["method"]))
    permutation = pd.read_csv(tables / "bonus_permutation_tests.csv")
    assert permutation["permutation_pvalue_two_sided"].between(0, 1).all()
    split_stability = pd.read_csv(tables / "bonus_split_stability.csv")
    assert split_stability["top20_technology_jaccard"].between(0, 1).all()
    umap_3d = pd.read_csv(tables / "bonus_umap_3d.csv")
    assert len(umap_3d) == bonus["umap_3d"]["sample_size"] == 10_000
    assert {"UMAP1", "UMAP2", "UMAP3", "cluster"}.issubset(umap_3d.columns)
    interactive_html = ROOT / bonus["umap_3d"]["html_path"]
    assert interactive_html.stat().st_size > 1_000_000
    assert "plotly" in interactive_html.read_text(encoding="utf-8")[:100_000].lower()

    assigner_path = artifacts / "models" / "phase3_consensus_assigner.joblib"
    joblib.load(assigner_path)
    assert _sha256(assigner_path) == registry["assigner_sha256"]

    notebook_results = {}
    for path in sorted((ROOT / "notebooks").glob("*_executed.ipynb")):
        notebook = nbformat.read(path, as_version=4)
        code_cells = [cell for cell in notebook.cells if cell.cell_type == "code"]
        errors = [
            output
            for cell in code_cells
            for output in cell.get("outputs", [])
            if output.output_type == "error"
        ]
        assert all(cell.get("execution_count") is not None for cell in code_cells)
        assert not errors
        notebook_results[path.name] = {"code_cells": len(code_cells), "errors": 0}

    report_hashes = {}
    for path in sorted(reports.glob("*Report_FA.docx")):
        document = Document(path)
        section = document.sections[0]
        assert abs(section.page_width - Cm(21.0)) < 10_000
        assert abs(section.page_height - Cm(29.7)) < 10_000
        assert section.different_first_page_header_footer
        assert document.styles["Normal"]._element.rPr.rFonts.get(qn("w:cs")) == "B Nazanin"
        # Desktop Word may rewrite built-in heading styles to theme references
        # while preserving the explicit font on every heading run.
        _assert_heading_fonts(document, "B Nazanin")
        assert all(table._tbl.tblPr.find(qn("w:bidiVisual")) is not None for table in document.tables)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        cover_text = "\n".join(paragraph.text for paragraph in document.paragraphs[:12])
        cover_text += "\n" + "\n".join(
            cell.text for table in document.tables[:2] for row in table.rows for cell in row.cells
        )
        expected_metadata = (
            REPORT_DETAILS["university"],
            REPORT_DETAILS["faculty"],
            REPORT_DETAILS["instructor"],
            *[name.strip() for name in REPORT_DETAILS["authors"].split("—")],
            *REPORT_DETAILS["student_ids"],
        )
        for expected in filter(None, expected_metadata):
            assert expected in cover_text, f"Missing cover metadata in {path.name}: {expected}"
        assert not any(marker in text for marker in ("Ø", "Ù", "Û", "â€"))
        assert "پس از بازکردن فایل" not in text
        report_hashes[path.name] = _sha256(path)
        pdf = path.with_suffix(".pdf")
        assert pdf.stat().st_size > 100_000
        assert pdf.read_bytes()[:4] == b"%PDF"
        with fitz.open(pdf) as pdf_document:
            assert len(pdf_document) > 0
            assert all(
                abs(page.rect.width - 595.32) < 1 and abs(page.rect.height - 841.92) < 1
                for page in pdf_document
            )
    assert report_hashes["Phase3_Report_FA.docx"] != report_hashes["Final_Report_FA.docx"]

    for path in (
        ROOT / "dashboard" / "app.py",
        ROOT / "scripts" / "run_all.py",
        ROOT / "src" / "stackoverflow_clustering" / "bonus.py",
    ):
        py_compile.compile(str(path), doraise=True)

    result = {
        "status": "passed",
        "rows_raw": 65_437,
        "rows_final": 60_023,
        "phase2_configurations": 78,
        "consensus_k": phase3["consensus_k"],
        "cluster_sizes": phase3["cluster_sizes"],
        "bonus_points_targeted": bonus["points_targeted"],
        "nmf_components": bonus["nmf"]["selected_components"],
        "umap_3d_sample_size": bonus["umap_3d"]["sample_size"],
        "notebooks": notebook_results,
        "reports": sorted(report_hashes),
    }
    (artifacts / "release_validation.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return result


if __name__ == "__main__":
    print(json.dumps(validate(), ensure_ascii=False, indent=2))
