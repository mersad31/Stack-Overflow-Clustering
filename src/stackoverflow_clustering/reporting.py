from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .config import project_path


NAVY = "203D68"
TEAL = "2D6F8C"
LIGHT_TEAL = "E8EEF4"
GOLD = "C7A64A"
LIGHT_GRAY = "F3F0EB"
BOX_FILL = "F7F4EF"
BORDER_BLUE = "8EA6C1"
WHITE = "FFFFFF"
BODY_FONT = "B Nazanin"
LATIN_FONT = BODY_FONT
HEADING_FONT = BODY_FONT
LOGO_PATH = Path(__file__).with_name("kntu_logo.png")

REPORT_DETAILS = {
    "university": "دانشگاه صنعتی خواجه نصیرالدین طوسی",
    "faculty": "دانشکده مهندسی کامپیوتر",
    "instructor": "دکتر پیشگو",
    "authors": "مرصاد احمدی — علی حسن زاده",
    "student_ids": ("40301274", "40304464"),
    "date": "تیر ۱۴۰۵ / July 2026",
}


def _set_paragraph_property(paragraph, name: str, value: str | None = None) -> None:
    if hasattr(paragraph, "_p"):
        p_pr = paragraph._p.get_or_add_pPr()
    else:
        p_pr = paragraph._element.get_or_add_pPr()
    element = p_pr.find(qn(f"w:{name}"))
    if element is None:
        element = OxmlElement(f"w:{name}")
        p_pr.append(element)
    if value is not None:
        element.set(qn("w:val"), value)


def _strip_arabic_diacritics(text: object) -> str:
    """Remove Arabic vowel and recitation marks without changing letters."""
    return "".join(
        character
        for character in str(text)
        if not (
            "\u0610" <= character <= "\u061a"
            or "\u064b" <= character <= "\u065f"
            or character == "\u0670"
            or "\u06d6" <= character <= "\u06ed"
        )
    )


def _rtl_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    # Word swaps visual left/right for bidi paragraphs. Using ``left`` here
    # produces a visually right-aligned Persian paragraph in exported PDFs.
    _set_paragraph_property(paragraph, "bidi")
    alignment_values = {
        WD_ALIGN_PARAGRAPH.LEFT: "right",
        WD_ALIGN_PARAGRAPH.RIGHT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
    }
    _set_paragraph_property(paragraph, "jc", alignment_values.get(alignment, "both"))
    _set_paragraph_property(paragraph, "widowControl")


def _rtl_run(run, font: str = BODY_FONT, size: float = 13, bold: bool | None = None) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT if font == BODY_FONT else font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT if font == BODY_FONT else font)
    run._element.rPr.rFonts.set(qn("w:cs"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    language = OxmlElement("w:lang")
    language.set(qn("w:bidi"), "fa-IR")
    language.set(qn("w:val"), "fa-IR")
    run._element.get_or_add_rPr().append(language)
    rtl = OxmlElement("w:rtl")
    run._element.get_or_add_rPr().append(rtl)


def _ltr_run(run, font: str = LATIN_FONT, size: float = 11, bold: bool | None = None) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    language = OxmlElement("w:lang")
    language.set(qn("w:val"), "en-US")
    run._element.get_or_add_rPr().append(language)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str = BORDER_BLUE, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def _paragraph_bottom_border(paragraph, color: str = NAVY, size: str = "10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)


def _set_cell_text(cell, value: object, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    text = "" if pd.isna(value) else _strip_arabic_diacritics(value)
    contains_persian = any("\u0600" <= character <= "\u06ff" for character in text)
    if contains_persian:
        _rtl_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        run = paragraph.add_run(text)
        _rtl_run(run, size=11.5, bold=bold)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        _ltr_run(run, size=10.5, bold=bold)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_property(paragraph, "bidi", "0")
    run = paragraph.add_run()
    _ltr_run(run, font=LATIN_FONT, size=9.5)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


class PersianReport:
    def __init__(self, title: str):
        self.document = Document()
        self.title = title
        self._configure_document()

    def _configure_document(self) -> None:
        doc = self.document
        section = doc.sections[0]
        section.top_margin = Cm(1.65)
        section.bottom_margin = Cm(1.65)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.orientation = WD_ORIENT.PORTRAIT
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)
        section.different_first_page_header_footer = True

        normal = doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(12.5)
        normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        normal._element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        _set_paragraph_property(normal, "bidi")
        _set_paragraph_property(normal, "jc", "both")

        for name, size, color in (
            ("Title", 24, NAVY),
            ("Heading 1", 16, NAVY),
            ("Heading 2", 13.5, NAVY),
            ("Heading 3", 12.5, TEAL),
        ):
            style = doc.styles[name]
            style.font.name = HEADING_FONT
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style._element.rPr.rFonts.set(qn("w:ascii"), HEADING_FONT)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), HEADING_FONT)
            style._element.rPr.rFonts.set(qn("w:cs"), HEADING_FONT)
            style.paragraph_format.space_before = Pt(13 if name == "Heading 1" else 8)
            style.paragraph_format.space_after = Pt(6 if name == "Heading 1" else 4)
            style.paragraph_format.keep_with_next = True
            _set_paragraph_property(style, "bidi")
            _set_paragraph_property(style, "jc", "left")
            _set_paragraph_property(style, "keepNext")

        for name in ("TOC 1", "TOC 2", "TOC 3"):
            if name not in doc.styles:
                continue
            style = doc.styles[name]
            style.font.name = BODY_FONT
            style.font.size = Pt(12.5 if name == "TOC 1" else 11.5)
            style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
            style._element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)
            style.paragraph_format.space_after = Pt(4)
            _set_paragraph_property(style, "bidi")
            _set_paragraph_property(style, "jc", "left")

        header = section.header.paragraphs[0]
        header.text = ""

        footer = section.footer.paragraphs[0]
        _page_number(footer)
        section.first_page_header.paragraphs[0].text = ""
        section.first_page_footer.paragraphs[0].text = ""

    def add_cover(
        self,
        subtitle: str,
        authors: str,
        date_text: str,
        university: str | None = None,
        faculty: str | None = None,
        instructor: str | None = None,
        student_ids: list[str] | None = None,
    ) -> None:
        doc = self.document
        logo_line = doc.add_paragraph()
        logo_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_line.paragraph_format.space_before = Pt(10)
        logo_line.paragraph_format.space_after = Pt(3)
        if LOGO_PATH.exists():
            logo_line.add_run().add_picture(str(LOGO_PATH), width=Inches(1.45))

        university_line = doc.add_paragraph()
        _rtl_paragraph(university_line, WD_ALIGN_PARAGRAPH.CENTER)
        university_line.paragraph_format.space_after = Pt(2)
        university_run = university_line.add_run(_strip_arabic_diacritics(university or "گزارش پروژه دانشگاهی"))
        _rtl_run(university_run, font=HEADING_FONT, size=14, bold=True)
        university_run.font.color.rgb = RGBColor.from_string(NAVY)
        if faculty:
            faculty_line = doc.add_paragraph()
            _rtl_paragraph(faculty_line, WD_ALIGN_PARAGRAPH.CENTER)
            faculty_line.paragraph_format.space_after = Pt(70)
            faculty_run = faculty_line.add_run(_strip_arabic_diacritics(faculty))
            _rtl_run(faculty_run, font=HEADING_FONT, size=11.5, bold=False)
            faculty_run.font.color.rgb = RGBColor.from_string(NAVY)
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(70)

        title_line = doc.add_paragraph()
        _rtl_paragraph(title_line, WD_ALIGN_PARAGRAPH.CENTER)
        title_line.paragraph_format.space_after = Pt(10)
        title_run = title_line.add_run(_strip_arabic_diacritics(subtitle))
        _rtl_run(title_run, font=HEADING_FONT, size=23, bold=True)
        title_run.font.color.rgb = RGBColor.from_string(NAVY)

        subject_line = doc.add_paragraph()
        _rtl_paragraph(subject_line, WD_ALIGN_PARAGRAPH.CENTER)
        subject_line.paragraph_format.space_after = Pt(34)
        subject_run = subject_line.add_run(_strip_arabic_diacritics(self.title))
        _rtl_run(subject_run, font=HEADING_FONT, size=14, bold=True)
        subject_run.font.color.rgb = RGBColor.from_string(TEAL)

        authors_line = doc.add_paragraph()
        _rtl_paragraph(authors_line, WD_ALIGN_PARAGRAPH.CENTER)
        authors_run = authors_line.add_run(_strip_arabic_diacritics(authors))
        _rtl_run(authors_run, font=HEADING_FONT, size=13.5, bold=True)
        authors_run.font.color.rgb = RGBColor.from_string(NAVY)

        if student_ids:
            id_line = doc.add_paragraph()
            id_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # The author line is RTL; reverse the LTR ID sequence so each ID
            # appears beneath its owner in the same visual order.
            id_run = id_line.add_run(" — ".join(reversed(student_ids)))
            _ltr_run(id_run, font=LATIN_FONT, size=11.5, bold=True)
            id_run.font.color.rgb = RGBColor.from_string(NAVY)

        if instructor:
            instructor_line = doc.add_paragraph()
            _rtl_paragraph(instructor_line, WD_ALIGN_PARAGRAPH.CENTER)
            instructor_run = instructor_line.add_run(_strip_arabic_diacritics(f"استاد: {instructor}"))
            _rtl_run(instructor_run, font=BODY_FONT, size=11.5)
            instructor_run.font.color.rgb = RGBColor.from_string(NAVY)

        date_line = doc.add_paragraph()
        _rtl_paragraph(date_line, WD_ALIGN_PARAGRAPH.CENTER)
        date_line.paragraph_format.space_before = Pt(15)
        date_run = date_line.add_run(_strip_arabic_diacritics(date_text))
        _rtl_run(date_run, font=HEADING_FONT, size=11.5)
        date_run.font.color.rgb = RGBColor.from_string(NAVY)
        doc.add_page_break()

    def add_toc(self) -> None:
        self.heading("فهرست مطالب", 1)
        paragraph = self.document.add_paragraph()
        _rtl_paragraph(paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "فهرست پس از بازکردن فایل در Word به‌روزرسانی می‌شود."
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, separate, placeholder, end])
        self.document.add_page_break()

    def heading(self, text: str, level: int = 1) -> None:
        paragraph = self.document.add_heading(_strip_arabic_diacritics(text), level=level)
        _rtl_paragraph(paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
        for run in paragraph.runs:
            _rtl_run(run, font=HEADING_FONT, size={1: 16, 2: 13.5, 3: 12.5}.get(level, 12), bold=True)
        if level == 1:
            _paragraph_bottom_border(paragraph)

    def paragraph(self, text: str, bold_lead: str | None = None) -> None:
        text = _strip_arabic_diacritics(text)
        bold_lead = _strip_arabic_diacritics(bold_lead) if bold_lead else None
        paragraph = self.document.add_paragraph()
        _rtl_paragraph(paragraph)
        if bold_lead and text.startswith(bold_lead):
            lead = paragraph.add_run(bold_lead)
            _rtl_run(lead, bold=True)
            body = paragraph.add_run(text[len(bold_lead) :])
            _rtl_run(body)
        else:
            run = paragraph.add_run(text)
            _rtl_run(run)

    def bullet(self, text: str) -> None:
        paragraph = self.document.add_paragraph(style="List Bullet")
        _rtl_paragraph(paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
        run = paragraph.add_run(_strip_arabic_diacritics(text))
        _rtl_run(run)

    def code_block(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_property(paragraph, "bidi", "0")
        _set_paragraph_property(paragraph, "jc", "left")
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.right_indent = Cm(0.5)
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(7)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), LIGHT_GRAY)
        paragraph._p.get_or_add_pPr().append(shading)
        run = paragraph.add_run(_strip_arabic_diacritics(text))
        _ltr_run(run, font=BODY_FONT, size=9.5)

    def callout(self, title: str, text: str) -> None:
        """Add a restrained academic note box matching the reference report."""
        table = self.document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        self._set_table_rtl(table)
        cell = table.cell(0, 0)
        _shade_cell(cell, BOX_FILL)
        _set_cell_border(cell)
        paragraph = cell.paragraphs[0]
        _rtl_paragraph(paragraph)
        title_run = paragraph.add_run(f"{_strip_arabic_diacritics(title)}  ")
        _rtl_run(title_run, font=HEADING_FONT, size=11.5, bold=True)
        title_run.font.color.rgb = RGBColor.from_string(NAVY)
        body_run = paragraph.add_run(_strip_arabic_diacritics(text))
        _rtl_run(body_run, size=11.5)
        self.document.add_paragraph().paragraph_format.space_after = Pt(1)

    def table(self, headers: Iterable[str], rows: Iterable[Iterable[object]], widths=None) -> None:
        headers = list(headers)
        rows = list(rows)
        table = self.document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        self._set_table_rtl(table)
        for idx, header in enumerate(headers):
            _shade_cell(table.cell(0, idx), NAVY)
            _set_cell_text(table.cell(0, idx), header, bold=True, color=WHITE)
        header_properties = table.rows[0]._tr.get_or_add_trPr()
        header_properties.append(OxmlElement("w:tblHeader"))
        # Keep the repeated header with at least the first data row.  Without
        # this, Word can leave a lone header row at the foot of a page.
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
                _set_paragraph_property(paragraph, "keepNext")
        for row_index, row in enumerate(rows):
            cells = table.add_row().cells
            table.rows[-1]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            for col_index, value in enumerate(row):
                _set_cell_text(cells[col_index], value)
                if row_index % 2:
                    _shade_cell(cells[col_index], LIGHT_GRAY)
        if widths:
            for row in table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = Cm(width)
        self.document.add_paragraph()

    def figure(self, path: Path, caption: str, width_inches: float = 6.5) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        _set_paragraph_property(paragraph, "keepNext")
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(width_inches))
        caption_p = self.document.add_paragraph()
        _rtl_paragraph(caption_p, WD_ALIGN_PARAGRAPH.CENTER)
        caption_p.paragraph_format.keep_together = True
        _set_paragraph_property(caption_p, "keepLines")
        caption_run = caption_p.add_run(_strip_arabic_diacritics(caption))
        _rtl_run(caption_run, font=HEADING_FONT, size=10.5, bold=True)
        caption_run.font.color.rgb = RGBColor.from_string(TEAL)

    @staticmethod
    def _set_table_rtl(table) -> None:
        table_pr = table._tbl.tblPr
        bidi = table_pr.find(qn("w:bidiVisual"))
        if bidi is None:
            table_pr.append(OxmlElement("w:bidiVisual"))
        margins = table_pr.find(qn("w:tblCellMar"))
        if margins is None:
            margins = OxmlElement("w:tblCellMar")
            for side, width in (("top", "90"), ("left", "110"), ("bottom", "90"), ("right", "110")):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), width)
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            table_pr.append(margins)

    def page_break(self) -> None:
        self.document.add_page_break()

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        # Enforce B Nazanin and remove Arabic combining marks everywhere,
        # including tables, headers, footers, and generated TOC text.
        for style in self.document.styles:
            if style._element.rPr is None:
                continue
            style.font.name = BODY_FONT
            for script in ("ascii", "hAnsi", "cs", "eastAsia"):
                style._element.rPr.rFonts.set(qn(f"w:{script}"), BODY_FONT)

        def normalize_paragraph(paragraph) -> None:
            for run in paragraph.runs:
                if run.text:
                    cleaned = _strip_arabic_diacritics(run.text)
                    if cleaned != run.text:
                        run.text = cleaned
                run.font.name = BODY_FONT
                properties = run._element.get_or_add_rPr()
                for script in ("ascii", "hAnsi", "cs", "eastAsia"):
                    properties.rFonts.set(qn(f"w:{script}"), BODY_FONT)

        def normalize_table(table) -> None:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        normalize_paragraph(paragraph)
                    for nested_table in cell.tables:
                        normalize_table(nested_table)

        for paragraph in self.document.paragraphs:
            normalize_paragraph(paragraph)
        for table in self.document.tables:
            normalize_table(table)
        for section in self.document.sections:
            for container in (
                section.header,
                section.footer,
                section.first_page_header,
                section.first_page_footer,
            ):
                for paragraph in container.paragraphs:
                    normalize_paragraph(paragraph)
                for table in container.tables:
                    normalize_table(table)
        settings = self.document.settings._element
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)
        self.document.save(path)


def add_project_cover(report: PersianReport, title: str) -> None:
    """Add the common course and author information to a report cover."""
    report.add_cover(
        title,
        REPORT_DETAILS["authors"],
        REPORT_DETAILS["date"],
        university=REPORT_DETAILS["university"],
        faculty=REPORT_DETAILS["faculty"],
        instructor=REPORT_DETAILS["instructor"],
        student_ids=list(REPORT_DETAILS["student_ids"]),
    )


def build_phase1_report(config: dict[str, Any]) -> Path:
    artifacts = project_path(config, "artifacts_dir")
    tables = project_path(config, "tables_dir")
    figures = project_path(config, "figures_dir")
    reports_dir = figures.parent

    ingestion = json.loads((artifacts / "ingestion_manifest.json").read_text(encoding="utf-8"))
    cleaning = json.loads((artifacts / "cleaning_summary.json").read_text(encoding="utf-8"))
    features = json.loads((artifacts / "feature_manifest.json").read_text(encoding="utf-8"))
    reduction = json.loads((artifacts / "reduction_manifest.json").read_text(encoding="utf-8"))
    decisions = pd.read_csv(tables / "cleaning_decision_log.csv")
    hopkins = pd.read_csv(tables / "hopkins_summary.csv")
    tech = pd.read_csv(tables / "technology_prevalence.csv").sort_values(
        "respondents", ascending=False
    )

    report = PersianReport("تحلیل خوشه‌ای پروفایل توسعه‌دهندگان")
    add_project_cover(report, "گزارش فاز اول: آماده‌سازی داده و ارزیابی گرایش به خوشه‌بندی")
    report.add_toc()

    report.heading("چکیده اجرایی", 1)
    report.paragraph(
        f"پیمایش Stack Overflow 2024 شامل {ingestion['rows']:,} پاسخ‌دهنده و {ingestion['source_columns']} ستون خام است. "
        f"پس از اعمال قواعد کیفیت و حذف {cleaning['all_technology_domains_missing_removed']:,} رکورد فاقد هرگونه شواهد پشته فناوری، "
        f"cohort نهایی {cleaning['rows_cleaned']:,} پاسخ‌دهنده شد. دو نمایش اصلی ساخته شد: X_full با "
        f"{features['representations']['X_full'][1]} ویژگی و X_tech_stack با {features['representations']['X_tech_stack'][1]} ویژگی. "
        f"پنجاه مؤلفه PCA مقدار {100*reduction['pca_cumulative_variance']:.1f}٪ واریانس را حفظ کرد. "
        "آزمون Hopkins در تمام نمایش‌ها بالاتر از ۰٫۷ بود و وجود ساختار غیرتصادفی را تأیید کرد؛ در عین حال VAT و UMAP ساختاری پیوسته، هم‌پوشان و دارای تغییر چگالی را نشان دادند. "
        "بنابراین فاز دوم باید روش‌های partitioning، سلسله‌مراتبی، چگالی‌محور، مدل‌مبنا و spectral را به‌صورت مقایسه‌ای بررسی کند."
    )
    report.table(
        ["شاخص کلیدی", "مقدار"],
        [
            ["پاسخ‌دهندگان خام", f"{ingestion['rows']:,}"],
            ["cohort پاک‌سازی‌شده", f"{cleaning['rows_cleaned']:,}"],
            ["ویژگی‌های X_full", features["representations"]["X_full"][1]],
            ["ویژگی‌های فناوری", features["technology_features"]],
            ["ویژگی‌های X_tech_stack", features["representations"]["X_tech_stack"][1]],
            ["واریانس حفظ‌شده در PCA-50", f"{100*reduction['pca_cumulative_variance']:.2f}٪"],
        ],
    )

    report.heading("۱. داده و خط لوله ورود", 1)
    report.paragraph(
        "داده اصلی survey_results_public.csv و فرهنگ داده survey_results_schema.csv هستند. pipeline پیش از خواندن داده، وجود تمام ستون‌های موردنیاز و schema فایل فرهنگ را کنترل می‌کند، فایل ۱۵۹ مگابایتی را در chunkهای ده‌هزارردیفی می‌خواند و نسخه منتخب را در قالب Parquet ذخیره می‌کند. برای داده خام، schema و خروجی میانی هش SHA-256 ثبت شده است تا هر اجرای آینده قابل ردیابی باشد."
    )
    report.paragraph(
        "از ۱۸ متغیر تحلیلی طرح اولیه، چهار متغیر DevType، EdLevel، AISelect و ConvertedCompYearly فقط برای تفسیر/تحلیل پایین‌دستی نگه داشته شدند. این جداسازی از ارزیابی دوری جلوگیری می‌کند. چهارده متغیر باقی‌مانده شامل تجربه، اندازه سازمان، وضعیت اشتغال، دورکاری، کشور و هفت حوزه فناوری هستند. ResponseId، MainBranch و Check نیز صرفا برای کنترل داده وارد شدند."
    )
    report.figure(
        figures / "phase1_missingness_selected_fields.png",
        "شکل ۱ — درصد داده گمشده در ستون‌های منتخب پیش از پاک‌سازی",
    )

    report.heading("۲. پاک‌سازی و مدیریت نوع داده", 1)
    report.paragraph(
        f"در داده هیچ شناسه تکراری و هیچ شکست صریح آزمون توجه وجود نداشت. تعداد {cleaning['rare_country_categories_grouped']} کشور با فراوانی کمتر از {cleaning['rare_country_threshold']} در گروه Other قرار گرفتند. "
        "رشته‌های تجربه به سال عددی تبدیل شدند؛ «کمتر از یک سال» برابر ۰٫۵ و «بیش از ۵۰ سال» برابر ۵۱ در نظر گرفته شد. Missing تجربه با میانه همان گروه کشور و سپس میانه سراسری تکمیل شد، اما indicator گمشدگی حفظ گردید. اندازه سازمان به midpoint بازه نگاشت و log1p شد. جبران خدمت فقط متغیر مکمل است و مقادیر غیرمثبت آن حذف می‌شوند."
    )
    important = decisions.sort_values("affected_rows", ascending=False).head(10)
    report.table(
        ["ستون", "مسئله", "اقدام", "ردیف متاثر"],
        important[["column", "issue", "action", "affected_rows"]].itertuples(index=False, name=None),
    )
    report.figure(
        figures / "phase1_experience_distributions.png",
        "شکل ۲ — توزیع سه سنجه تجربه پیش از imputation",
    )

    report.heading("۳. مهندسی ویژگی", 1)
    report.paragraph(
        "هفت ستون فناوری با جداکننده semicolon شکافته و بدون حذف هیچ فناوری به ۲۴۰ ویژگی multi-hot تبدیل شدند. یک ویژگی TechnologyBreadth مجموع فناوری‌های پاسخ‌دهنده را ثبت می‌کند؛ بنابراین X_tech_stack دقیقا ۲۴۱ ویژگی دارد. Employment نیز به ۹ ویژگی multi-hot تبدیل شد. Country و RemoteWork به one-hot تبدیل شدند."
    )
    report.paragraph(
        "ویژگی‌های عددی مشتق‌شده شامل سه تجربه تکمیل‌شده، میانه توافقی تجربه، اختلاف سنجه‌های تجربه، نسبت تجربه حرفه‌ای به کل، log اندازه سازمان، indicatorهای گمشدگی و breadth هر حوزه فناوری است. ترکیب این بلوک با متغیرهای دودویی، X_full با ۳۲۸ ویژگی را ساخته است."
    )
    inventory = pd.read_csv(tables / "feature_inventory.csv")
    report.table(
        ["بلوک ویژگی", "تعداد"],
        inventory[["block", "features"]].itertuples(index=False, name=None),
    )
    report.figure(
        figures / "phase1_top_technologies.png",
        "شکل ۳ — ۲۵ فناوری پرتکرار در cohort نهایی",
    )

    report.heading("۴. مقیاس‌بندی و انتخاب scaler", 1)
    report.paragraph(
        "StandardScaler، RobustScaler و MinMaxScaler روی بلوک عددی مقایسه شدند؛ ویژگی‌های one-hot/multi-hot به‌صورت دودویی باقی ماندند. نمودار قبل/بعد نشان می‌دهد هر سه روش هندسه را به‌شدت تغییر می‌دهند. داده تجربه و اندازه سازمان دارای دنباله‌های بلند و پرت‌های واقعی است؛ بنابراین RobustScaler با center میانه و IQR به‌عنوان انتخاب اصلی قفل شد."
    )
    scaling_rows = hopkins[hopkins["representation"].str.startswith("X_full_")].copy()
    scaling_rows["mean"] = scaling_rows["mean"].map(lambda x: f"{x:.4f}")
    scaling_rows["std"] = scaling_rows["std"].map(lambda x: f"{x:.4f}")
    report.table(
        ["نمایش", "میانگین Hopkins", "انحراف معیار"],
        scaling_rows[["representation", "mean", "std"]].itertuples(index=False, name=None),
    )
    report.paragraph(
        "StandardScaler مقدار Hopkins بالاتری دارد، اما Hopkins معیار کیفیت خوشه‌بندی یا تابع هدف انتخاب scaler نیست؛ فقط فاصله از تصادفی‌بودن را می‌سنجد و به تغییر هندسه حساس است. انتخاب Robust بر مقاومت در برابر tailهای تجربه/اندازه سازمان مبتنی است. در فاز سوم حساسیت clustering نهایی به هر سه scaler با ARI گزارش خواهد شد."
    )
    report.figure(
        figures / "phase1_scaler_comparison.png",
        "شکل ۴ — هندسه تجربه و اندازه سازمان پیش و پس از سه روش مقیاس‌بندی",
        width_inches=6.8,
    )

    report.heading("۵. کاهش بعد", 1)
    report.paragraph(
        f"PCA روی X_full مبتنی بر RobustScaler اجرا شد. پنجاه مؤلفه {100*reduction['pca_cumulative_variance']:.2f}٪ واریانس را حفظ کردند؛ ۸ مؤلفه نخست از ۵۰٪ و ۳۰ مؤلفه از ۷۰٪ عبور می‌کنند. "
        "انتخاب ۵۰ مؤلفه با طرح اولیه سازگار است و مصالحه‌ای میان حفظ اطلاعات و کاهش تمرکز فاصله‌ها فراهم می‌کند. برای ماتریس sparse فناوری نیز Truncated SVD پنجاه‌بعدی، "
        f"{100*reduction['svd_cumulative_variance']:.2f}٪ واریانس را نگه داشت. UMAP روی نمونه ثابت {reduction['umap_sample_size']:,}تایی برای تشخیص کیفی ساختار اجرا شد و ورودی نهایی الگوریتم‌های فاصله‌محور نیست."
    )
    report.figure(
        figures / "phase1_pca_explained_variance.png",
        "شکل ۵ — منحنی واریانس تجمعی PCA",
    )
    report.figure(
        figures / "phase1_pca_density.png",
        "شکل ۶ — چگالی پاسخ‌دهندگان در دو مؤلفه نخست PCA",
    )
    report.figure(
        figures / "phase1_umap_density.png",
        "شکل ۷ — نمایش غیرخطی UMAP روی نمونه ثابت ۱۵هزارتایی",
    )

    report.heading("۶. معیار فاصله و شباهت", 1)
    report.paragraph(
        "فاصله Euclidean در فضای PCA مقیاس‌شده معیار اصلی K-Means، Ward و GMM است، زیرا مؤلفه‌ها متعامد و عددی‌اند. "
        "Manhattan و Gower در طراحی اولیه به‌عنوان گزینه‌های جایگزین بررسی شدند، اما مقایسه نهایی خانواده‌ها برای حفظ هندسه و نمونه مشترک "
        "بر Euclidean در فضای PCA ثابت شد. Jaccard نیز فقط در سنجش پایداری واژگان فناوری میان دو split تجربه به‌کار رفت."
    )

    report.heading("۷. ارزیابی گرایش به خوشه‌بندی", 1)
    report.paragraph(
        "Hopkins در ده تکرار با seedهای ثبت‌شده و نمونه ۱۵۰۰تایی محاسبه شد. تمام میانگین‌ها از ۰٫۷ بیشترند و به‌ویژه PCA گرایش قوی ۰٫۸۵۵ را نشان می‌دهد. انحراف معیارهای نزدیک ۰٫۰۰۱ نشان می‌دهد نتیجه به نمونه تصادفی خاص وابسته نیست."
    )
    hopkins_display = hopkins.copy()
    for column in ["mean", "std", "min", "max"]:
        hopkins_display[column] = hopkins_display[column].map(lambda x: f"{x:.4f}")
    report.table(
        ["نمایش", "میانگین", "انحراف معیار", "کمینه", "بیشینه"],
        hopkins_display[["representation", "mean", "std", "min", "max"]].itertuples(index=False, name=None),
    )
    report.figure(
        figures / "phase1_vat_heatmap.png",
        "شکل ۸ — ماتریس عدم‌شباهت بازمرتب‌شده با VAT روی نمونه هزارنفری PCA",
    )
    report.paragraph(
        "VAT بلوک‌های کاملا جدا و مربعی نشان نمی‌دهد؛ به‌جای آن گرادیان‌ها و نواحی پیوسته دیده می‌شود. این نتیجه با UMAP سازگار است: ساختار واقعی وجود دارد، اما احتمالا شامل overlap، manifold و تغییر چگالی است. بنابراین نباید یک پاسخ K-Means را بدون مقایسه HDBSCAN، Spectral، GMM و روش سلسله‌مراتبی به‌عنوان حقیقت طبیعی پذیرفت."
    )

    report.heading("۸. EDA خوشه‌محور", 1)
    top_rows = tech.head(12).copy()
    top_rows["pct_of_cohort"] = top_rows["pct_of_cohort"].map(lambda x: f"{x:.2f}٪")
    report.table(
        ["حوزه", "فناوری", "پاسخ‌دهنده", "درصد cohort"],
        top_rows[["domain", "technology", "respondents", "pct_of_cohort"]].itertuples(index=False, name=None),
    )
    report.paragraph(
        "JavaScript، HTML/CSS، Python و SQL هر یک در بیش از نیمی از cohort دیده می‌شوند؛ بنابراین به‌تنهایی جداکننده کافی نیستند. ترکیب فناوری‌ها، breadth، تجربه و context کاری برای ساخت persona اهمیت بیشتری دارد. ماتریس Spearman و Cramér’s V نیز برای کنترل هم‌خطی و ارتباط متغیرهای رده‌ای تولید و در بسته artifact ذخیره شده است."
    )
    report.figure(
        figures / "phase1_numeric_correlation.png",
        "شکل ۹ — همبستگی Spearman ویژگی‌های عددی مهندسی‌شده",
    )
    report.figure(
        figures / "phase1_categorical_cramers_v.png",
        "شکل ۱۰ — ارتباط رده‌ای با Cramér’s V اصلاح‌شده",
    )

    report.heading("۹. محدودیت‌ها و تصمیم ورود به فاز دوم", 1)
    report.bullet("گمشدگی WorkExp و برخی حوزه‌های فناوری بالاست؛ indicator گمشدگی حفظ شده، اما imputation همچنان عدم‌قطعیت ایجاد می‌کند.")
    report.bullet("Hopkins قوی وجود ساختار غیرتصادفی را نشان می‌دهد، نه تعداد یا معنای خوشه‌ها را.")
    report.bullet("PCA-50 حدود ۲۲٫۲٪ واریانس را حذف می‌کند؛ نتایج بر X_full و X_tech_stack نیز برای حساسیت مقایسه می‌شوند.")
    report.bullet("DevType، EdLevel و AISelect برچسب مرجع نیستند و فقط متغیرهای کمکی برای تفسیرند؛ معیارهای خارجی با این محدودیت گزارش می‌شوند.")
    report.paragraph(
        "نتایج فاز اول ادامه تحلیل خوشه‌ای را توجیه می‌کند. در فاز دوم k با Elbow/Kneedle، Silhouette، Gap، DB/CH/Dunn و BIC بررسی می‌شود و پایداری seed/bootstrap همراه با توافق الگوریتم‌ها سنجیده خواهد شد."
    )

    report.heading("پیوست الف — ثبت کامل تصمیم‌های پاک‌سازی", 1)
    report.table(
        ["ستون", "مسئله", "اقدام", "توجیه", "ردیف متاثر"],
        decisions[["column", "issue", "action", "rationale", "affected_rows"]].itertuples(index=False, name=None),
    )
    report.heading("پیوست ب — بازتولیدپذیری و اخلاق", 1)
    report.paragraph(
        "تمام روش‌های تصادفی seed پایه 42 دارند. مسیرها و پارامترها از config.yaml خوانده می‌شوند و اجرای کامل فاز با فرمان زیر امکان‌پذیر است:"
    )
    report.code_block(".\\venv\\Scripts\\python.exe -m scripts.run_phase1 --config config.yaml")
    report.paragraph(
        "منبع داده Stack Overflow Developer Survey 2024 است. داده فقط برای تحلیل دانشگاهی استفاده شده و هیچ شناسایی فردی انجام نشده است."
    )

    output = reports_dir / "Phase1_Report_FA.docx"
    report.save(output)
    return output
